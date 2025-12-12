"""Word document viewer."""

from textual import work
from textual.app import ComposeResult
from textual.containers import VerticalScroll
from textual.widgets import Static

from quickview.viewers.base import BaseViewer
from quickview.utils import escape_markup


class DocxWidget(Static):
    """Widget to display Word document content."""

    def __init__(self, filepath: str, **kwargs):
        super().__init__(**kwargs)
        self.filepath = filepath

    def on_mount(self) -> None:
        self._load_docx()

    @work(thread=True)
    def _load_docx(self) -> None:
        try:
            from docx import Document
        except ImportError:
            self.app.call_from_thread(
                self.update,
                "[red]Error: python-docx not installed. Run: pip install quickview[docx][/red]",
            )
            return

        try:
            doc = Document(self.filepath)
            lines = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    if para.style.name.startswith("Heading"):
                        level = para.style.name[-1] if para.style.name[-1].isdigit() else "1"
                        lines.append(f"[bold cyan]{'#' * int(level)} {text}[/bold cyan]")
                    else:
                        text = escape_markup(text)
                        lines.append(text)
                    lines.append("")

            for table in doc.tables:
                lines.append("[dim]─" * 40 + "[/dim]")
                for row in table.rows:
                    cells = [escape_markup(cell.text.strip()) for cell in row.cells]
                    lines.append(" │ ".join(cells))
                lines.append("[dim]─" * 40 + "[/dim]")
                lines.append("")

            content = "\n".join(lines) if lines else "[dim]<Empty document>[/dim]"
            self.app.call_from_thread(self.update, content)

        except Exception as e:
            self.app.call_from_thread(
                self.update,
                f"[red]Error loading Word document: {e}[/red]",
            )


class DocxViewer(BaseViewer):
    """Viewer for Word documents."""

    extensions = [".docx"]

    def compose(self) -> ComposeResult:
        with VerticalScroll():
            yield DocxWidget(self.filepath, id="docx-content")

    def load(self) -> None:
        # Loading is handled by DocxWidget.on_mount
        pass
